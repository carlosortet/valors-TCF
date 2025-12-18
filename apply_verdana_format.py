#!/usr/bin/env python3
"""
Script para aplicar formato Verdana a un documento Word.
Basado en el formato especificado en format_verdana.xml
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def apply_verdana_format(input_file, output_file):
    """Aplica formato Verdana tamaño 10 a todo el documento."""

    print(f"Abriendo documento: {input_file}")
    doc = Document(input_file)

    # Aplicar formato a todos los párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Verdana'
            run.font.size = Pt(10)  # 20 half-points = 10 points
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro

    # Aplicar formato a tablas si existen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Verdana'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # Ajustar estilos del documento
    styles = doc.styles

    # Estilo Normal
    if 'Normal' in styles:
        style = styles['Normal']
        style.font.name = 'Verdana'
        style.font.size = Pt(10)

    # Estilos de títulos
    heading_sizes = {
        'Heading 1': 16,
        'Heading 2': 14,
        'Heading 3': 12,
        'Heading 4': 11,
        'Heading 5': 10,
        'Heading 6': 10,
    }

    for heading_name, size in heading_sizes.items():
        if heading_name in styles:
            style = styles[heading_name]
            style.font.name = 'Verdana'
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

    print(f"Guardando documento con formato aplicado: {output_file}")
    doc.save(output_file)
    print("✓ Formato Verdana aplicado correctamente")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python3 apply_verdana_format.py <input.docx> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        apply_verdana_format(input_file, output_file)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
